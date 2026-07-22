#!/usr/bin/env python3
"""
MdToDocx.py - Combine markdown files into a single searchable .docx

Scans a source folder for *.md files (all markdown, not just synopses) and
combines them into one Word document with preserved formatting (headings,
lists, bold/italic, tables, code blocks, hyperlinks, images). A clickable Table
of Contents is inserted at the top.

Images: standard markdown ![alt](path) is embedded. Relative paths resolve
against the folder of the markdown file that references them, so you keep
images in (e.g.) Documentation/images/ and never paste them in by hand. Remote
http(s) images are linked rather than embedded; missing files degrade to a
visible note instead of failing the run.

Image sizing: every image renders at DEFAULT_IMAGE_WIDTH_IN for a consistent
look; images narrower than that keep their native size. Override one image with
a width in the markdown title slot: ![alt](path "w=4") or "w=320px".

Minimal tkinter GUI: source folder picker, output folder picker, Run button,
log tail.

Dependencies:
    pip install python-docx markdown-it-py

Logs to: C:/Logs/MdToDocx.log
Output:  MdCombined_YYYY-MM-DD_HHMMSS.docx  (in the chosen output folder)

Usage:
    python MdToDocx.py                # launches the GUI
"""

import logging
import os
import re
import sys
import threading
import tkinter as tk
from datetime import datetime
from pathlib import Path
from tkinter import filedialog, messagebox, scrolledtext, ttk

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from markdown_it import MarkdownIt

# =============================================================================
# CONFIG
# =============================================================================

LOG_DIR = Path("C:/Logs")
LOG_FILE = LOG_DIR / "MdToDocx.log"

# Glob pattern(s) to scan for. Accepts all .md files by default.
# To restrict to specific suffixes, add entries like "*_Synopsis.md".
DEFAULT_GLOBS = ["*.md"]

# Folders to skip when walking the source tree
SKIP_FOLDERS = {".idea", ".git", ".venv", "__pycache__", "node_modules",
                "CommitsGH", "PunchlistReview", "Archive"}

DOC_TITLE = "Markdown Documents"

# Max embedded image width, in inches. Letter page with 1" margins gives ~6.5"
# of usable width; 6.0 leaves a little breathing room. Images wider than this
# are scaled down proportionally; smaller images keep their native size (so
# icons/small diagrams are never upscaled).
MAX_IMAGE_WIDTH_IN = 6.0

# Default render width, in inches, applied to every embedded image so they
# come out a consistent size (good for screenshots). Images NARROWER than this
# keep their native size, so small icons aren't upscaled. Set to None to fall
# back to "native size, capped at MAX_IMAGE_WIDTH_IN".
# Per-image override: put a width in the markdown title slot, e.g.
#   ![alt](./Images/foo.png "w=4")      -> 4 inches
#   ![alt](./Images/foo.png "w=320px")  -> 320 px (at 96 dpi)
# An explicit per-image width always wins, even if it upscales.
DEFAULT_IMAGE_WIDTH_IN = 6.0

# =============================================================================
# LOGGING
# =============================================================================

def setup_logging():
    """Set up file + console logging. No date suffix on the log file name."""
    LOG_DIR.mkdir(parents=True, exist_ok=True)
    logger_ = logging.getLogger("MdToDocx")
    logger_.setLevel(logging.INFO)
    logger_.handlers.clear()

    fmt = logging.Formatter("%(asctime)s [%(levelname)s] %(message)s",
                            datefmt="%Y-%m-%d %H:%M:%S")

    fh = logging.FileHandler(LOG_FILE, encoding="utf-8")
    fh.setFormatter(fmt)
    logger_.addHandler(fh)

    ch = logging.StreamHandler(sys.stdout)
    ch.setFormatter(fmt)
    logger_.addHandler(ch)

    return logger_


logger = setup_logging()


# =============================================================================
# FILE DISCOVERY
# =============================================================================

def find_markdown_files(source_dir: Path, patterns=None):
    """
    Look in source_dir (top level only) for files matching any of the glob
    patterns. Default is all *.md files. Returns a sorted list.
    """
    if patterns is None:
        patterns = DEFAULT_GLOBS

    found = []
    for p in source_dir.iterdir():
        if not p.is_file():
            continue
        if p.suffix.lower() != ".md":
            continue
        name_lower = p.name.lower()
        for pattern in patterns:
            # "*.md" matches everything; "*_Synopsis.md" matches that suffix
            suffix = pattern.replace("*", "").lower()
            if name_lower.endswith(suffix):
                if p not in found:
                    found.append(p)
                break

    found.sort(key=lambda x: x.name.lower())
    return found


# =============================================================================
# DOCX STYLE SETUP
# =============================================================================

def configure_styles(doc: Document):
    """Configure built-in styles so headings/body are consistent and TOC works."""
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    heading_sizes = {
        "Heading 1": 20,
        "Heading 2": 16,
        "Heading 3": 13,
        "Heading 4": 12,
        "Heading 5": 11,
        "Heading 6": 11,
    }
    for name, size in heading_sizes.items():
        try:
            s = styles[name]
            s.font.name = "Calibri"
            s.font.size = Pt(size)
            s.font.bold = True
            s.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)  # CRP navy
        except KeyError:
            pass

    # Code Block paragraph style — create if it doesn't exist
    existing_names = [s.name for s in styles]
    if "Code Block" not in existing_names:
        try:
            code_style = styles.add_style("Code Block", 1)  # 1 = paragraph style
            code_style.font.name = "Consolas"
            code_style.font.size = Pt(9)
            code_style.base_style = styles["Normal"]
        except Exception:
            pass


def add_page_break(doc: Document):
    """Add a hard page break."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_toc(doc: Document):
    """
    Insert a Word TOC field. It will render empty until the user opens the doc
    and updates fields (F9 or 'Yes' to the update prompt).
    """
    p = doc.add_paragraph()
    run = p.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    # \o "1-4" = include headings 1-4; \h = hyperlinks; \z = hide tab leader in web view; \u = use outline levels
    instr_text.text = 'TOC \\o "1-4" \\h \\z \\u'

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click here and choose 'Update Field' to populate the Table of Contents."

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    r_element = run._r
    r_element.append(fld_char_begin)
    r_element.append(instr_text)
    r_element.append(fld_char_separate)
    r_element.append(placeholder)
    r_element.append(fld_char_end)


def add_hyperlink(paragraph, url, text):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)

    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_cell_shading(cell, hex_color):
    """Apply background shading to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


# =============================================================================
# MARKDOWN -> DOCX RENDERING
# =============================================================================

class MarkdownToDocx:
    """
    Walk markdown-it tokens and emit docx paragraphs/runs/tables.

    Heading level is shifted by heading_offset so file contents nest under the
    file's Heading 1 title (i.e. # in markdown becomes Heading 2 in the docx).
    """

    def __init__(self, doc: Document, heading_offset: int = 1):
        self.doc = doc
        self.heading_offset = heading_offset
        # Folder the current markdown file lives in; relative image paths
        # resolve against this. Set per-file in render().
        self.base_dir = None
        self.md = (MarkdownIt("commonmark", {"html": False})
                   .enable("table")
                   .enable("strikethrough"))

    # --- public entry point --------------------------------------------------

    def render(self, markdown_text: str, base_dir: Path = None):
        self.base_dir = base_dir
        tokens = self.md.parse(markdown_text)
        self._render_tokens(tokens)

    # --- top-level token walker ---------------------------------------------

    def _render_tokens(self, tokens):
        i = 0
        while i < len(tokens):
            tok = tokens[i]
            t = tok.type

            if t == "heading_open":
                level = int(tok.tag[1])
                shifted = min(level + self.heading_offset, 6)
                style_name = f"Heading {shifted}"
                inline = tokens[i + 1]
                para = self.doc.add_paragraph(style=style_name)
                self._render_inline(inline.children or [], para)
                while tokens[i].type != "heading_close":
                    i += 1
                i += 1
                continue

            if t == "paragraph_open":
                inline = tokens[i + 1]
                para = self.doc.add_paragraph()
                self._render_inline(inline.children or [], para)
                while tokens[i].type != "paragraph_close":
                    i += 1
                i += 1
                continue

            if t == "bullet_list_open" or t == "ordered_list_open":
                ordered = (t == "ordered_list_open")
                close_type = "ordered_list_close" if ordered else "bullet_list_close"
                end_idx = self._find_matching_close(tokens, i, t, close_type)
                self._render_list(tokens[i + 1:end_idx], ordered=ordered, level=0)
                i = end_idx + 1
                continue

            if t == "fence" or t == "code_block":
                self._render_code_block(tok.content)
                i += 1
                continue

            if t == "blockquote_open":
                end_idx = self._find_matching_close(
                    tokens, i, "blockquote_open", "blockquote_close")
                self._render_blockquote(tokens[i + 1:end_idx])
                i = end_idx + 1
                continue

            if t == "hr":
                p = self.doc.add_paragraph()
                p_pr = p._p.get_or_add_pPr()
                p_bdr = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "6")
                bottom.set(qn("w:space"), "1")
                bottom.set(qn("w:color"), "808080")
                p_bdr.append(bottom)
                p_pr.append(p_bdr)
                i += 1
                continue

            if t == "table_open":
                end_idx = self._find_matching_close(
                    tokens, i, "table_open", "table_close")
                self._render_table(tokens[i:end_idx + 1])
                i = end_idx + 1
                continue

            i += 1

    def _find_matching_close(self, tokens, start_idx, open_type, close_type):
        """Find the index of the close token matching tokens[start_idx]."""
        depth = 1
        i = start_idx + 1
        while i < len(tokens):
            if tokens[i].type == open_type:
                depth += 1
            elif tokens[i].type == close_type:
                depth -= 1
                if depth == 0:
                    return i
            i += 1
        return len(tokens) - 1

    # --- inline rendering ----------------------------------------------------

    def _render_inline(self, children, paragraph, base_state=None):
        """Render inline tokens into runs on the given paragraph."""
        state = dict(base_state or {"bold": False, "italic": False, "strike": False})
        i = 0
        while i < len(children):
            c = children[i]
            t = c.type

            if t == "text":
                run = paragraph.add_run(c.content)
                self._apply_state(run, state)
            elif t == "strong_open":
                state["bold"] = True
            elif t == "strong_close":
                state["bold"] = False
            elif t == "em_open":
                state["italic"] = True
            elif t == "em_close":
                state["italic"] = False
            elif t == "s_open":
                state["strike"] = True
            elif t == "s_close":
                state["strike"] = False
            elif t == "code_inline":
                run = paragraph.add_run(c.content)
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                r_pr = run._r.get_or_add_rPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F2F2F2")
                r_pr.append(shd)
            elif t == "image":
                self._render_image(c, paragraph)
            elif t == "softbreak":
                paragraph.add_run(" ")
            elif t == "hardbreak":
                paragraph.add_run().add_break()
            elif t == "link_open":
                href = dict(c.attrs or {}).get("href", "")
                link_text_parts = []
                j = i + 1
                while j < len(children) and children[j].type != "link_close":
                    if children[j].type == "text":
                        link_text_parts.append(children[j].content)
                    elif children[j].type == "code_inline":
                        link_text_parts.append(children[j].content)
                    j += 1
                link_text = "".join(link_text_parts) or href
                if href:
                    add_hyperlink(paragraph, href, link_text)
                else:
                    paragraph.add_run(link_text)
                i = j  # will be incremented to skip link_close
            elif t == "link_close":
                pass

            i += 1

    def _apply_state(self, run, state):
        if state.get("bold"):
            run.bold = True
        if state.get("italic"):
            run.italic = True
        if state.get("strike"):
            run.font.strike = True

    # --- images --------------------------------------------------------------

    def _render_image(self, token, paragraph):
        """
        Embed an image referenced by ![alt](src). Relative paths resolve
        against the current markdown file's folder (self.base_dir). Remote
        URLs are linked rather than embedded. Missing/unreadable files
        degrade to a visible note instead of crashing the build.
        """
        attrs = dict(token.attrs or {})
        src = (attrs.get("src") or "").strip()
        alt = (token.content or attrs.get("alt") or "image").strip()

        if not src:
            return

        # Remote images can't be embedded offline — link them instead.
        if src.lower().startswith(("http://", "https://")):
            add_hyperlink(paragraph, src, f"[Image: {alt}]")
            logger.warning(f"  Remote image linked, not embedded: {src}")
            return

        img_path = Path(src)
        if not img_path.is_absolute():
            base = self.base_dir or Path.cwd()
            img_path = (base / img_path).resolve()

        if not img_path.is_file():
            run = paragraph.add_run(f"[Image not found: {src}]")
            run.italic = True
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            logger.warning(f"  Image not found: {img_path}")
            return

        try:
            width = self._resolve_image_width(token, img_path)
            run = paragraph.add_run()
            if width is not None:
                run.add_picture(str(img_path), width=width)
            else:
                run.add_picture(str(img_path))
            logger.info(f"  Embedded image: {img_path.name}")
        except Exception as e:
            run = paragraph.add_run(f"[Image failed: {src}]")
            run.italic = True
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            logger.warning(f"  Failed to embed {img_path}: {e}")

    def _resolve_image_width(self, token, img_path: Path):
        """
        Decide the render width (an Inches value, or None = keep native size).
        Precedence:
          1. Explicit per-image hint in the markdown title  ->  always wins.
          2. DEFAULT_IMAGE_WIDTH_IN, if set  ->  uniform width for images that
             are wider than it; narrower images stay native (no upscaling).
          3. Hard cap at MAX_IMAGE_WIDTH_IN.
        """
        # 1. per-image override: ![alt](path "w=4")  /  "w=320px"
        hint = self._parse_width_hint(dict(token.attrs or {}).get("title", ""))
        if hint is not None:
            return hint

        native_in = self._native_width_in(img_path)  # None if unreadable

        # 2. uniform default
        if DEFAULT_IMAGE_WIDTH_IN is not None:
            if native_in is None or native_in > DEFAULT_IMAGE_WIDTH_IN:
                return Inches(DEFAULT_IMAGE_WIDTH_IN)
            return None  # smaller than the default — leave it alone

        # 3. cap only
        if native_in is None or native_in > MAX_IMAGE_WIDTH_IN:
            return Inches(MAX_IMAGE_WIDTH_IN)
        return None

    @staticmethod
    def _native_width_in(img_path: Path):
        """Native image width in inches, or None if it can't be read.
        Uses python-docx's own image reader — no extra dependency (e.g. PIL)."""
        try:
            from docx.image.image import Image as _DocxImage
            img = _DocxImage.from_file(str(img_path))
            return img.px_width / (img.horz_dpi or 96)
        except Exception:
            return None

    @staticmethod
    def _parse_width_hint(title: str):
        """Parse a width from the markdown title slot into an Inches value.
        Accepts: 'w=4', 'width=4', '4in', '4"', 'w=320px', '320px'. px at 96 dpi.
        Returns None if no width is present."""
        if not title:
            return None
        t = title.lower()
        m = re.search(r'(?:w|width)\s*=\s*(\d+(?:\.\d+)?)\s*(px|in|")?', t)
        if not m:
            m = re.search(r'(\d+(?:\.\d+)?)\s*(px|in|")', t)
        if not m:
            return None
        val = float(m.group(1))
        unit = (m.group(2) or "in").replace('"', "in")
        return Inches(val / 96.0) if unit == "px" else Inches(val)

    # --- lists ---------------------------------------------------------------

    def _render_list(self, body_tokens, ordered: bool, level: int):
        """
        Render the contents of a list (between list_open and list_close).
        body_tokens is the slice of tokens inside the list.
        """
        base_style = "List Number" if ordered else "List Bullet"

        i = 0
        while i < len(body_tokens):
            tok = body_tokens[i]
            if tok.type == "list_item_open":
                end_idx = self._find_matching_close(
                    body_tokens, i, "list_item_open", "list_item_close")
                self._render_list_item(body_tokens[i + 1:end_idx], base_style, level)
                i = end_idx + 1
                continue
            i += 1

    def _render_list_item(self, item_tokens, base_style, level):
        """
        Render a single list item. The first paragraph-like block gets the
        bullet/number style; subsequent blocks render inline with indentation.
        """
        # Word ships List Bullet, List Bullet 2 ... List Bullet 5 (same for Number)
        if level == 0:
            style_name = base_style
        else:
            style_name = f"{base_style} {min(level + 1, 5)}"

        first_para_done = False
        i = 0
        while i < len(item_tokens):
            tok = item_tokens[i]

            if tok.type == "paragraph_open":
                inline = item_tokens[i + 1]
                if not first_para_done:
                    try:
                        para = self.doc.add_paragraph(style=style_name)
                    except KeyError:
                        para = self.doc.add_paragraph(style=base_style)
                    first_para_done = True
                else:
                    para = self.doc.add_paragraph()
                    para.paragraph_format.left_indent = Inches(0.5 * (level + 1))
                self._render_inline(inline.children or [], para)
                # Skip to paragraph_close
                while i < len(item_tokens) and item_tokens[i].type != "paragraph_close":
                    i += 1
                i += 1
                continue

            if tok.type in ("bullet_list_open", "ordered_list_open"):
                ordered = (tok.type == "ordered_list_open")
                close_type = "ordered_list_close" if ordered else "bullet_list_close"
                end_idx = self._find_matching_close(item_tokens, i, tok.type, close_type)
                self._render_list(item_tokens[i + 1:end_idx],
                                  ordered=ordered, level=level + 1)
                i = end_idx + 1
                continue

            if tok.type in ("fence", "code_block"):
                self._render_code_block(tok.content)
                i += 1
                continue

            i += 1

    # --- code blocks ---------------------------------------------------------

    def _render_code_block(self, content: str):
        """Render a fenced or indented code block as a styled paragraph."""
        content = content.rstrip("\n")
        try:
            para = self.doc.add_paragraph(style="Code Block")
        except KeyError:
            para = self.doc.add_paragraph()

        p_pr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F5F5F5")
        p_pr.append(shd)

        lines = content.split("\n")
        for idx, line in enumerate(lines):
            run = para.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            if idx < len(lines) - 1:
                run.add_break()

    # --- blockquotes ---------------------------------------------------------

    def _render_blockquote(self, inner_tokens):
        """Render blockquote contents as quote-styled paragraphs."""
        start_index = len(self.doc.paragraphs)
        self._render_tokens(inner_tokens)
        for p in self.doc.paragraphs[start_index:]:
            try:
                p.style = self.doc.styles["Quote"]
            except KeyError:
                p.paragraph_format.left_indent = Inches(0.5)
                for run in p.runs:
                    run.italic = True

    # --- tables --------------------------------------------------------------

    def _render_table(self, table_tokens):
        """
        Render a GFM pipe table.
        table_tokens should include table_open ... table_close.
        """
        rows = []  # [(is_header, [cell_inline_children, ...]), ...]
        i = 0
        while i < len(table_tokens):
            tok = table_tokens[i]
            if tok.type == "tr_open":
                is_header = False
                cells = []
                j = i + 1
                while j < len(table_tokens) and table_tokens[j].type != "tr_close":
                    if table_tokens[j].type in ("th_open", "td_open"):
                        if table_tokens[j].type == "th_open":
                            is_header = True
                        inline_tok = (table_tokens[j + 1]
                                      if j + 1 < len(table_tokens) else None)
                        cell_inline = (inline_tok.children
                                       if inline_tok and inline_tok.children else [])
                        cells.append(cell_inline)
                        while (j < len(table_tokens)
                               and table_tokens[j].type not in ("th_close", "td_close")):
                            j += 1
                    j += 1
                rows.append((is_header, cells))
                i = j + 1
                continue
            i += 1

        if not rows:
            return

        ncols = max(len(r[1]) for r in rows)
        table = self.doc.add_table(rows=len(rows), cols=ncols)
        try:
            table.style = "Light Grid Accent 1"
        except KeyError:
            pass
        table.autofit = True

        for r_idx, (is_header, cells) in enumerate(rows):
            row = table.rows[r_idx]
            for c_idx in range(ncols):
                cell = row.cells[c_idx]
                cell.text = ""
                para = cell.paragraphs[0]
                if c_idx < len(cells):
                    self._render_inline(cells[c_idx], para)

                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

                if is_header:
                    set_cell_shading(cell, "1A3A5C")
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Spacer paragraph after the table
        self.doc.add_paragraph()


# =============================================================================
# DOCUMENT ASSEMBLY
# =============================================================================

def title_from_filename(path: Path) -> str:
    """Use the file stem as the section title, stripping common suffixes."""
    stem = path.stem
    # Strip known suffixes for cleaner titles; add more as needed
    for suffix in ("_Synopsis", "_DecisionLog", "_FINAL"):
        stem = re.sub(rf"{suffix}$", "", stem, flags=re.IGNORECASE)
    return stem


def build_document(md_files, output_path: Path):
    """Build the .docx. Single-file mode skips the cover page and TOC."""
    doc = Document()
    configure_styles(doc)

    single = (len(md_files) == 1)

    if not single:
        # --- Cover page (multi-file only) ---
        cover = doc.add_paragraph()
        cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_run = cover.add_run(DOC_TITLE)
        cover_run.font.size = Pt(28)
        cover_run.bold = True
        cover_run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subtitle.add_run(
            f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        sub_run.italic = True
        sub_run.font.size = Pt(12)

        doc.add_paragraph()
        included_header = doc.add_paragraph()
        ih_run = included_header.add_run(f"Files included ({len(md_files)}):")
        ih_run.bold = True
        for f in md_files:
            doc.add_paragraph(f.name, style="List Bullet")

        add_page_break(doc)

        # --- Table of Contents ---
        doc.add_paragraph("Table of Contents", style="Heading 1")
        add_toc(doc)
        add_page_break(doc)

    # --- Render content ---
    # Single file: headings stay at their native level (offset=0)
    # Multi file: headings shift down one level under the file-title Heading 1
    renderer = MarkdownToDocx(doc, heading_offset=0 if single else 1)

    for idx, md_file in enumerate(md_files):
        logger.info(f"Processing: {md_file.name}")
        try:
            text = md_file.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            text = md_file.read_text(encoding="latin-1")
            logger.warning(f"  {md_file.name} was not UTF-8 - read as latin-1")

        if not single:
            title = title_from_filename(md_file)
            doc.add_paragraph(title, style="Heading 1")

        renderer.render(text, base_dir=md_file.parent)

        if idx < len(md_files) - 1:
            add_page_break(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info(f"Wrote: {output_path}")


# =============================================================================
# GUI
# =============================================================================

class MdToDocxGUI:
    def __init__(self, root):
        self.root = root
        root.title("Markdown to DOCX")
        root.geometry("720x560")

        frm = ttk.Frame(root, padding=10)
        frm.pack(fill=tk.BOTH, expand=True)

        # Mode selector
        self.mode_var = tk.StringVar(value="single")
        mode_frame = ttk.LabelFrame(frm, text="Mode", padding=5)
        mode_frame.grid(row=0, column=0, columnspan=2, sticky="we", pady=(0, 10))
        ttk.Radiobutton(mode_frame, text="Single .md file",
                        variable=self.mode_var, value="single",
                        command=self._update_mode).pack(side=tk.LEFT, padx=(0, 20))
        ttk.Radiobutton(mode_frame, text="All .md files in a folder",
                        variable=self.mode_var, value="folder",
                        command=self._update_mode).pack(side=tk.LEFT)

        # Source — single file
        self.single_label = ttk.Label(frm, text="Markdown file:")
        self.single_label.grid(row=1, column=0, sticky="w")
        self.file_var = tk.StringVar()
        self.file_entry = ttk.Entry(frm, textvariable=self.file_var, width=70)
        self.file_entry.grid(row=2, column=0, sticky="we", padx=(0, 5))
        self.file_btn = ttk.Button(frm, text="Browse...", command=self.pick_file)
        self.file_btn.grid(row=2, column=1, sticky="e")

        # Source — folder
        self.folder_label = ttk.Label(frm, text="Source folder (contains *.md files):")
        self.folder_var = tk.StringVar()
        self.folder_entry = ttk.Entry(frm, textvariable=self.folder_var, width=70)
        self.folder_btn = ttk.Button(frm, text="Browse...", command=self.pick_source)

        # Output folder
        ttk.Label(frm, text="Output folder:").grid(
            row=5, column=0, sticky="w", pady=(10, 0))
        self.output_var = tk.StringVar()
        ttk.Entry(frm, textvariable=self.output_var, width=70).grid(
            row=6, column=0, sticky="we", padx=(0, 5))
        ttk.Button(frm, text="Browse...", command=self.pick_output).grid(
            row=6, column=1, sticky="e")

        # Action buttons
        btn_frame = ttk.Frame(frm)
        btn_frame.grid(row=7, column=0, columnspan=2, pady=15, sticky="w")
        self.run_btn = ttk.Button(btn_frame, text="Generate DOCX",
                                  command=self.run_generation)
        self.run_btn.pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(btn_frame, text="Open Log",
                   command=self.open_log).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(btn_frame, text="Close",
                   command=root.destroy).pack(side=tk.LEFT)

        # Log / status area
        ttk.Label(frm, text="Status:").grid(row=8, column=0, sticky="w")
        self.log_widget = scrolledtext.ScrolledText(
            frm, height=16, wrap=tk.WORD, font=("Consolas", 9))
        self.log_widget.grid(row=9, column=0, columnspan=2, sticky="nsew", pady=(5, 0))

        frm.columnconfigure(0, weight=1)
        frm.rowconfigure(9, weight=1)

        self._install_gui_log_handler()
        self._update_mode()
        self._log("Ready. Pick a file or folder to begin.")

    # --- mode toggle ---

    def _update_mode(self):
        """Show/hide the appropriate source picker based on mode."""
        if self.mode_var.get() == "single":
            # Show file picker, hide folder picker
            self.single_label.grid(row=1, column=0, sticky="w")
            self.file_entry.grid(row=2, column=0, sticky="we", padx=(0, 5))
            self.file_btn.grid(row=2, column=1, sticky="e")
            self.folder_label.grid_remove()
            self.folder_entry.grid_remove()
            self.folder_btn.grid_remove()
        else:
            # Show folder picker, hide file picker
            self.single_label.grid_remove()
            self.file_entry.grid_remove()
            self.file_btn.grid_remove()
            self.folder_label.grid(row=3, column=0, sticky="w")
            self.folder_entry.grid(row=4, column=0, sticky="we", padx=(0, 5))
            self.folder_btn.grid(row=4, column=1, sticky="e")

    # --- UI helpers ---

    def _install_gui_log_handler(self):
        class GuiHandler(logging.Handler):
            def __init__(self, widget):
                super().__init__()
                self.widget = widget

            def emit(self, record):
                msg = self.format(record)
                try:
                    self.widget.after(0, self.widget.insert, tk.END, msg + "\n")
                    self.widget.after(0, self.widget.see, tk.END)
                except Exception:
                    pass

        handler = GuiHandler(self.log_widget)
        handler.setFormatter(logging.Formatter("%(asctime)s  %(message)s",
                                               datefmt="%H:%M:%S"))
        logger.addHandler(handler)

    def _log(self, msg):
        self.log_widget.insert(tk.END, msg + "\n")
        self.log_widget.see(tk.END)

    def pick_file(self):
        filepath = filedialog.askopenfilename(
            title="Select a Markdown file",
            filetypes=[("Markdown files", "*.md"), ("All files", "*.*")])
        if filepath:
            self.file_var.set(filepath)
            if not self.output_var.get():
                self.output_var.set(str(Path(filepath).parent))

    def pick_source(self):
        folder = filedialog.askdirectory(title="Select source folder")
        if folder:
            self.folder_var.set(folder)
            if not self.output_var.get():
                self.output_var.set(folder)

    def pick_output(self):
        folder = filedialog.askdirectory(title="Select output folder")
        if folder:
            self.output_var.set(folder)

    def open_log(self):
        if LOG_FILE.exists():
            os.startfile(str(LOG_FILE))
        else:
            messagebox.showinfo("Log", f"No log file at {LOG_FILE} yet.")

    # --- run ---

    def run_generation(self):
        mode = self.mode_var.get()
        output = self.output_var.get().strip()

        if mode == "single":
            source = self.file_var.get().strip()
            src_path = Path(source) if source else None
            if not src_path or not src_path.is_file():
                messagebox.showerror("Source file",
                                     "Pick a valid .md file.")
                return
            if not output:
                output = str(src_path.parent)
                self.output_var.set(output)
        else:
            source = self.folder_var.get().strip()
            if not source or not Path(source).is_dir():
                messagebox.showerror("Source folder",
                                     "Pick a valid source folder.")
                return
            if not output:
                output = source
                self.output_var.set(output)

        self.run_btn.config(state=tk.DISABLED)
        t = threading.Thread(target=self._run_worker,
                             args=(mode, source, Path(output)),
                             daemon=True)
        t.start()

    def _run_worker(self, mode: str, source: str, output: Path):
        try:
            if mode == "single":
                src_path = Path(source)
                files = [src_path]
                logger.info(f"Single file: {src_path.name}")
                stem = title_from_filename(src_path)
                out_path = output / f"{stem}.docx"
            else:
                src_dir = Path(source)
                logger.info(f"Scanning: {src_dir}")
                files = find_markdown_files(src_dir)
                if not files:
                    logger.warning(f"No *.md files found in {src_dir}")
                    messagebox.showwarning(
                        "No files",
                        f"No *.md files found in:\n{src_dir}")
                    return

                logger.info(f"Found {len(files)} file(s)")
                for f in files:
                    logger.info(f"  - {f.name}")

                ts = datetime.now().strftime("%Y-%m-%d_%H%M%S")
                out_path = output / f"MdCombined_{ts}.docx"

            build_document(files, out_path)

            logger.info("Done.")
            if messagebox.askyesno(
                    "Complete",
                    f"Generated:\n{out_path}\n\nOpen it now?"):
                os.startfile(str(out_path))
        except Exception as e:
            logger.exception(f"Failed: {e}")
            messagebox.showerror("Error", f"Generation failed:\n{e}")
        finally:
            self.run_btn.config(state=tk.NORMAL)


# =============================================================================
# MAIN
# =============================================================================

def main():
    """No arguments -> the tkinter GUI (unchanged). Any of --file/--source-folder
    -> headless conversion, so the tool can be scripted from a .bat / pipeline."""
    import argparse
    parser = argparse.ArgumentParser(
        description="Combine markdown into a .docx. With no arguments, launches the GUI.")
    parser.add_argument("--file", help="a single .md file -> <stem>.docx")
    parser.add_argument("--source-folder", help="a folder of .md files -> one combined .docx")
    parser.add_argument("--out-dir", help="output folder (default: alongside the source)")
    parser.add_argument("--out-name", help="output filename for --file mode (default: <stem>.docx)")
    args = parser.parse_args()

    if not args.file and not args.source_folder:
        root = tk.Tk()
        MdToDocxGUI(root)
        root.mainloop()
        return

    if args.file:
        src = Path(args.file)
        if not src.is_file():
            logger.error("Not a file: %s", src)
            raise SystemExit(2)
        files = [src]
        out_dir = Path(args.out_dir) if args.out_dir else src.parent
        name = args.out_name or f"{title_from_filename(src)}.docx"
        if not name.lower().endswith(".docx"):
            name += ".docx"
        out_path = out_dir / name
    else:
        src_dir = Path(args.source_folder)
        if not src_dir.is_dir():
            logger.error("Not a folder: %s", src_dir)
            raise SystemExit(2)
        files = find_markdown_files(src_dir)
        if not files:
            logger.error("No .md files found in %s", src_dir)
            raise SystemExit(1)
        out_dir = Path(args.out_dir) if args.out_dir else src_dir
        ts = datetime.now().strftime("%Y-%m-%d_%H%M%S")
        out_path = out_dir / f"MdCombined_{ts}.docx"

    build_document(files, out_path)
    print(f"Wrote: {out_path}")


if __name__ == "__main__":
    main()