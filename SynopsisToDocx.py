#!/usr/bin/env python3
"""
SynopsisToDocx.py - Combine project synopsis markdown files into a single searchable .docx

Scans a source folder for *_Synopsis.md files and combines them into one Word
document with preserved formatting (headings, lists, bold/italic, tables, code blocks,
hyperlinks). A clickable Table of Contents is inserted at the top.

Minimal tkinter GUI: source folder picker, output folder picker, Run button, log tail.

Dependencies:
    pip install python-docx markdown-it-py

Logs to: C:/Logs/SynopsisToDocx.log
Output:  Synopses_YYYY-MM-DD_HHMMSS.docx  (in the chosen output folder)

Usage:
    python SynopsisToDocx.py          # launches the GUI
"""

import logging
import os
import re
import sys
import threading
import tkinter as tk
from datetime import datetime
from pathlib import Path
from tkinter import filedialog, messagebox, scrolledtext, ttk

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from markdown_it import MarkdownIt

# =============================================================================
# CONFIG
# =============================================================================

LOG_DIR = Path("C:/Logs")
LOG_FILE = LOG_DIR / "SynopsisToDocx.log"

# Glob pattern(s) to scan for. Structured as a list so adding *_DecisionLog.md
# (or similar) later is a one-line change.
DEFAULT_GLOBS = ["*_Synopsis.md"]

# Folders to skip when walking the source tree
SKIP_FOLDERS = {".idea", ".git", ".venv", "__pycache__", "node_modules",
                "CommitsGH", "PunchlistReview", "Archive"}

DOC_TITLE = "CRP Project Synopses"

# =============================================================================
# LOGGING
# =============================================================================

def setup_logging():
    """Set up file + console logging. No date suffix on the log file name."""
    LOG_DIR.mkdir(parents=True, exist_ok=True)
    logger_ = logging.getLogger("SynopsisToDocx")
    logger_.setLevel(logging.INFO)
    logger_.handlers.clear()

    fmt = logging.Formatter("%(asctime)s [%(levelname)s] %(message)s",
                            datefmt="%Y-%m-%d %H:%M:%S")

    fh = logging.FileHandler(LOG_FILE, encoding="utf-8")
    fh.setFormatter(fmt)
    logger_.addHandler(fh)

    ch = logging.StreamHandler(sys.stdout)
    ch.setFormatter(fmt)
    logger_.addHandler(ch)

    return logger_


logger = setup_logging()


# =============================================================================
# FILE DISCOVERY
# =============================================================================

def find_markdown_files(source_dir: Path, patterns=None):
    """
    Look in source_dir (top level only — synopsis files live directly in it)
    for files matching any of the glob patterns. Returns a sorted list.
    """
    if patterns is None:
        patterns = DEFAULT_GLOBS

    found = []
    for p in source_dir.iterdir():
        if not p.is_file():
            continue
        if p.suffix.lower() != ".md":
            continue
        name_lower = p.name.lower()
        for pattern in patterns:
            suffix = pattern.replace("*", "").lower()
            if name_lower.endswith(suffix):
                if p not in found:
                    found.append(p)
                break

    found.sort(key=lambda x: x.name.lower())
    return found


# =============================================================================
# DOCX STYLE SETUP
# =============================================================================

def configure_styles(doc: Document):
    """Configure built-in styles so headings/body are consistent and TOC works."""
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    heading_sizes = {
        "Heading 1": 20,
        "Heading 2": 16,
        "Heading 3": 13,
        "Heading 4": 12,
        "Heading 5": 11,
        "Heading 6": 11,
    }
    for name, size in heading_sizes.items():
        try:
            s = styles[name]
            s.font.name = "Calibri"
            s.font.size = Pt(size)
            s.font.bold = True
            s.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)  # CRP navy
        except KeyError:
            pass

    # Code Block paragraph style — create if it doesn't exist
    existing_names = [s.name for s in styles]
    if "Code Block" not in existing_names:
        try:
            code_style = styles.add_style("Code Block", 1)  # 1 = paragraph style
            code_style.font.name = "Consolas"
            code_style.font.size = Pt(9)
            code_style.base_style = styles["Normal"]
        except Exception:
            pass


def add_page_break(doc: Document):
    """Add a hard page break."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_toc(doc: Document):
    """
    Insert a Word TOC field. It will render empty until the user opens the doc
    and updates fields (F9 or 'Yes' to the update prompt).
    """
    p = doc.add_paragraph()
    run = p.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    # \o "1-4" = include headings 1-4; \h = hyperlinks; \z = hide tab leader in web view; \u = use outline levels
    instr_text.text = 'TOC \\o "1-4" \\h \\z \\u'

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click here and choose 'Update Field' to populate the Table of Contents."

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    r_element = run._r
    r_element.append(fld_char_begin)
    r_element.append(instr_text)
    r_element.append(fld_char_separate)
    r_element.append(placeholder)
    r_element.append(fld_char_end)


def add_hyperlink(paragraph, url, text):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)

    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_cell_shading(cell, hex_color):
    """Apply background shading to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


# =============================================================================
# MARKDOWN -> DOCX RENDERING
# =============================================================================

class MarkdownToDocx:
    """
    Walk markdown-it tokens and emit docx paragraphs/runs/tables.

    Heading level is shifted by heading_offset so file contents nest under the
    file's Heading 1 title (i.e. # in markdown becomes Heading 2 in the docx).
    """

    def __init__(self, doc: Document, heading_offset: int = 1):
        self.doc = doc
        self.heading_offset = heading_offset
        self.md = (MarkdownIt("commonmark", {"html": False})
                   .enable("table")
                   .enable("strikethrough"))

    # --- public entry point --------------------------------------------------

    def render(self, markdown_text: str):
        tokens = self.md.parse(markdown_text)
        self._render_tokens(tokens)

    # --- top-level token walker ---------------------------------------------

    def _render_tokens(self, tokens):
        i = 0
        while i < len(tokens):
            tok = tokens[i]
            t = tok.type

            if t == "heading_open":
                level = int(tok.tag[1])
                shifted = min(level + self.heading_offset, 6)
                style_name = f"Heading {shifted}"
                inline = tokens[i + 1]
                para = self.doc.add_paragraph(style=style_name)
                self._render_inline(inline.children or [], para)
                while tokens[i].type != "heading_close":
                    i += 1
                i += 1
                continue

            if t == "paragraph_open":
                inline = tokens[i + 1]
                para = self.doc.add_paragraph()
                self._render_inline(inline.children or [], para)
                while tokens[i].type != "paragraph_close":
                    i += 1
                i += 1
                continue

            if t == "bullet_list_open" or t == "ordered_list_open":
                ordered = (t == "ordered_list_open")
                close_type = "ordered_list_close" if ordered else "bullet_list_close"
                end_idx = self._find_matching_close(tokens, i, t, close_type)
                self._render_list(tokens[i + 1:end_idx], ordered=ordered, level=0)
                i = end_idx + 1
                continue

            if t == "fence" or t == "code_block":
                self._render_code_block(tok.content)
                i += 1
                continue

            if t == "blockquote_open":
                end_idx = self._find_matching_close(
                    tokens, i, "blockquote_open", "blockquote_close")
                self._render_blockquote(tokens[i + 1:end_idx])
                i = end_idx + 1
                continue

            if t == "hr":
                p = self.doc.add_paragraph()
                p_pr = p._p.get_or_add_pPr()
                p_bdr = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "6")
                bottom.set(qn("w:space"), "1")
                bottom.set(qn("w:color"), "808080")
                p_bdr.append(bottom)
                p_pr.append(p_bdr)
                i += 1
                continue

            if t == "table_open":
                end_idx = self._find_matching_close(
                    tokens, i, "table_open", "table_close")
                self._render_table(tokens[i:end_idx + 1])
                i = end_idx + 1
                continue

            i += 1

    def _find_matching_close(self, tokens, start_idx, open_type, close_type):
        """Find the index of the close token matching tokens[start_idx]."""
        depth = 1
        i = start_idx + 1
        while i < len(tokens):
            if tokens[i].type == open_type:
                depth += 1
            elif tokens[i].type == close_type:
                depth -= 1
                if depth == 0:
                    return i
            i += 1
        return len(tokens) - 1

    # --- inline rendering ----------------------------------------------------

    def _render_inline(self, children, paragraph, base_state=None):
        """Render inline tokens into runs on the given paragraph."""
        state = dict(base_state or {"bold": False, "italic": False, "strike": False})
        i = 0
        while i < len(children):
            c = children[i]
            t = c.type

            if t == "text":
                run = paragraph.add_run(c.content)
                self._apply_state(run, state)
            elif t == "strong_open":
                state["bold"] = True
            elif t == "strong_close":
                state["bold"] = False
            elif t == "em_open":
                state["italic"] = True
            elif t == "em_close":
                state["italic"] = False
            elif t == "s_open":
                state["strike"] = True
            elif t == "s_close":
                state["strike"] = False
            elif t == "code_inline":
                run = paragraph.add_run(c.content)
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                r_pr = run._r.get_or_add_rPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F2F2F2")
                r_pr.append(shd)
            elif t == "softbreak":
                paragraph.add_run(" ")
            elif t == "hardbreak":
                paragraph.add_run().add_break()
            elif t == "link_open":
                href = dict(c.attrs or {}).get("href", "")
                link_text_parts = []
                j = i + 1
                while j < len(children) and children[j].type != "link_close":
                    if children[j].type == "text":
                        link_text_parts.append(children[j].content)
                    elif children[j].type == "code_inline":
                        link_text_parts.append(children[j].content)
                    j += 1
                link_text = "".join(link_text_parts) or href
                if href:
                    add_hyperlink(paragraph, href, link_text)
                else:
                    paragraph.add_run(link_text)
                i = j  # will be incremented to skip link_close
            elif t == "link_close":
                pass

            i += 1

    def _apply_state(self, run, state):
        if state.get("bold"):
            run.bold = True
        if state.get("italic"):
            run.italic = True
        if state.get("strike"):
            run.font.strike = True

    # --- lists ---------------------------------------------------------------

    def _render_list(self, body_tokens, ordered: bool, level: int):
        """
        Render the contents of a list (between list_open and list_close).
        body_tokens is the slice of tokens inside the list.
        """
        base_style = "List Number" if ordered else "List Bullet"

        i = 0
        while i < len(body_tokens):
            tok = body_tokens[i]
            if tok.type == "list_item_open":
                end_idx = self._find_matching_close(
                    body_tokens, i, "list_item_open", "list_item_close")
                self._render_list_item(body_tokens[i + 1:end_idx], base_style, level)
                i = end_idx + 1
                continue
            i += 1

    def _render_list_item(self, item_tokens, base_style, level):
        """
        Render a single list item. The first paragraph-like block gets the
        bullet/number style; subsequent blocks render inline with indentation.
        """
        # Word ships List Bullet, List Bullet 2 ... List Bullet 5 (same for Number)
        if level == 0:
            style_name = base_style
        else:
            style_name = f"{base_style} {min(level + 1, 5)}"

        first_para_done = False
        i = 0
        while i < len(item_tokens):
            tok = item_tokens[i]

            if tok.type == "paragraph_open":
                inline = item_tokens[i + 1]
                if not first_para_done:
                    try:
                        para = self.doc.add_paragraph(style=style_name)
                    except KeyError:
                        para = self.doc.add_paragraph(style=base_style)
                    first_para_done = True
                else:
                    para = self.doc.add_paragraph()
                    para.paragraph_format.left_indent = Inches(0.5 * (level + 1))
                self._render_inline(inline.children or [], para)
                # Skip to paragraph_close
                while i < len(item_tokens) and item_tokens[i].type != "paragraph_close":
                    i += 1
                i += 1
                continue

            if tok.type in ("bullet_list_open", "ordered_list_open"):
                ordered = (tok.type == "ordered_list_open")
                close_type = "ordered_list_close" if ordered else "bullet_list_close"
                end_idx = self._find_matching_close(item_tokens, i, tok.type, close_type)
                self._render_list(item_tokens[i + 1:end_idx],
                                  ordered=ordered, level=level + 1)
                i = end_idx + 1
                continue

            if tok.type in ("fence", "code_block"):
                self._render_code_block(tok.content)
                i += 1
                continue

            i += 1

    # --- code blocks ---------------------------------------------------------

    def _render_code_block(self, content: str):
        """Render a fenced or indented code block as a styled paragraph."""
        content = content.rstrip("\n")
        try:
            para = self.doc.add_paragraph(style="Code Block")
        except KeyError:
            para = self.doc.add_paragraph()

        p_pr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F5F5F5")
        p_pr.append(shd)

        lines = content.split("\n")
        for idx, line in enumerate(lines):
            run = para.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            if idx < len(lines) - 1:
                run.add_break()

    # --- blockquotes ---------------------------------------------------------

    def _render_blockquote(self, inner_tokens):
        """Render blockquote contents as quote-styled paragraphs."""
        start_index = len(self.doc.paragraphs)
        self._render_tokens(inner_tokens)
        for p in self.doc.paragraphs[start_index:]:
            try:
                p.style = self.doc.styles["Quote"]
            except KeyError:
                p.paragraph_format.left_indent = Inches(0.5)
                for run in p.runs:
                    run.italic = True

    # --- tables --------------------------------------------------------------

    def _render_table(self, table_tokens):
        """
        Render a GFM pipe table.
        table_tokens should include table_open ... table_close.
        """
        rows = []  # [(is_header, [cell_inline_children, ...]), ...]
        i = 0
        while i < len(table_tokens):
            tok = table_tokens[i]
            if tok.type == "tr_open":
                is_header = False
                cells = []
                j = i + 1
                while j < len(table_tokens) and table_tokens[j].type != "tr_close":
                    if table_tokens[j].type in ("th_open", "td_open"):
                        if table_tokens[j].type == "th_open":
                            is_header = True
                        inline_tok = (table_tokens[j + 1]
                                      if j + 1 < len(table_tokens) else None)
                        cell_inline = (inline_tok.children
                                       if inline_tok and inline_tok.children else [])
                        cells.append(cell_inline)
                        while (j < len(table_tokens)
                               and table_tokens[j].type not in ("th_close", "td_close")):
                            j += 1
                    j += 1
                rows.append((is_header, cells))
                i = j + 1
                continue
            i += 1

        if not rows:
            return

        ncols = max(len(r[1]) for r in rows)
        table = self.doc.add_table(rows=len(rows), cols=ncols)
        try:
            table.style = "Light Grid Accent 1"
        except KeyError:
            pass
        table.autofit = True

        for r_idx, (is_header, cells) in enumerate(rows):
            row = table.rows[r_idx]
            for c_idx in range(ncols):
                cell = row.cells[c_idx]
                cell.text = ""
                para = cell.paragraphs[0]
                if c_idx < len(cells):
                    self._render_inline(cells[c_idx], para)

                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

                if is_header:
                    set_cell_shading(cell, "1A3A5C")
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Spacer paragraph after the table
        self.doc.add_paragraph()


# =============================================================================
# DOCUMENT ASSEMBLY
# =============================================================================

def title_from_filename(path: Path) -> str:
    """Strip _Synopsis (or _DecisionLog) from the stem for display."""
    stem = path.stem
    stem = re.sub(r"_Synopsis$", "", stem, flags=re.IGNORECASE)
    stem = re.sub(r"_DecisionLog$", "", stem, flags=re.IGNORECASE)
    return stem


def build_document(md_files, output_path: Path):
    """Build the combined .docx."""
    doc = Document()
    configure_styles(doc)

    # --- Cover page ---
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_run = cover.add_run(DOC_TITLE)
    cover_run.font.size = Pt(28)
    cover_run.bold = True
    cover_run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(
        f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    sub_run.italic = True
    sub_run.font.size = Pt(12)

    doc.add_paragraph()
    included_header = doc.add_paragraph()
    ih_run = included_header.add_run(f"Files included ({len(md_files)}):")
    ih_run.bold = True
    for f in md_files:
        doc.add_paragraph(f.name, style="List Bullet")

    add_page_break(doc)

    # --- Table of Contents ---
    doc.add_paragraph("Table of Contents", style="Heading 1")
    add_toc(doc)
    add_page_break(doc)

    # --- One section per file ---
    renderer = MarkdownToDocx(doc, heading_offset=1)

    for idx, md_file in enumerate(md_files):
        logger.info(f"Processing: {md_file.name}")
        try:
            text = md_file.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            text = md_file.read_text(encoding="latin-1")
            logger.warning(f"  {md_file.name} was not UTF-8 - read as latin-1")

        title = title_from_filename(md_file)
        doc.add_paragraph(title, style="Heading 1")

        renderer.render(text)

        if idx < len(md_files) - 1:
            add_page_break(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info(f"Wrote: {output_path}")


# =============================================================================
# GUI
# =============================================================================

class SynopsisToDocxGUI:
    def __init__(self, root):
        self.root = root
        root.title("Synopsis to DOCX")
        root.geometry("720x520")

        frm = ttk.Frame(root, padding=10)
        frm.pack(fill=tk.BOTH, expand=True)

        # Source folder
        ttk.Label(frm, text="Source folder (contains *_Synopsis.md):").grid(
            row=0, column=0, sticky="w")
        self.source_var = tk.StringVar()
        ttk.Entry(frm, textvariable=self.source_var, width=70).grid(
            row=1, column=0, sticky="we", padx=(0, 5))
        ttk.Button(frm, text="Browse...", command=self.pick_source).grid(
            row=1, column=1, sticky="e")

        # Output folder
        ttk.Label(frm, text="Output folder:").grid(
            row=2, column=0, sticky="w", pady=(10, 0))
        self.output_var = tk.StringVar()
        ttk.Entry(frm, textvariable=self.output_var, width=70).grid(
            row=3, column=0, sticky="we", padx=(0, 5))
        ttk.Button(frm, text="Browse...", command=self.pick_output).grid(
            row=3, column=1, sticky="e")

        # Action buttons
        btn_frame = ttk.Frame(frm)
        btn_frame.grid(row=4, column=0, columnspan=2, pady=15, sticky="w")
        self.run_btn = ttk.Button(btn_frame, text="Generate DOCX",
                                  command=self.run_generation)
        self.run_btn.pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(btn_frame, text="Open Log",
                   command=self.open_log).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(btn_frame, text="Close",
                   command=root.destroy).pack(side=tk.LEFT)

        # Log / status area
        ttk.Label(frm, text="Status:").grid(row=5, column=0, sticky="w")
        self.log_widget = scrolledtext.ScrolledText(
            frm, height=18, wrap=tk.WORD, font=("Consolas", 9))
        self.log_widget.grid(row=6, column=0, columnspan=2, sticky="nsew", pady=(5, 0))

        frm.columnconfigure(0, weight=1)
        frm.rowconfigure(6, weight=1)

        self._install_gui_log_handler()
        self._log("Ready. Pick a source folder to begin.")

    # --- UI helpers ---

    def _install_gui_log_handler(self):
        class GuiHandler(logging.Handler):
            def __init__(self, widget):
                super().__init__()
                self.widget = widget

            def emit(self, record):
                msg = self.format(record)
                try:
                    self.widget.after(0, self.widget.insert, tk.END, msg + "\n")
                    self.widget.after(0, self.widget.see, tk.END)
                except Exception:
                    pass

        handler = GuiHandler(self.log_widget)
        handler.setFormatter(logging.Formatter("%(asctime)s  %(message)s",
                                               datefmt="%H:%M:%S"))
        logger.addHandler(handler)

    def _log(self, msg):
        self.log_widget.insert(tk.END, msg + "\n")
        self.log_widget.see(tk.END)

    def pick_source(self):
        folder = filedialog.askdirectory(title="Select source folder")
        if folder:
            self.source_var.set(folder)
            if not self.output_var.get():
                self.output_var.set(folder)

    def pick_output(self):
        folder = filedialog.askdirectory(title="Select output folder")
        if folder:
            self.output_var.set(folder)

    def open_log(self):
        if LOG_FILE.exists():
            os.startfile(str(LOG_FILE))
        else:
            messagebox.showinfo("Log", f"No log file at {LOG_FILE} yet.")

    # --- run ---

    def run_generation(self):
        source = self.source_var.get().strip()
        output = self.output_var.get().strip()
        if not source or not Path(source).is_dir():
            messagebox.showerror("Source folder", "Pick a valid source folder.")
            return
        if not output:
            output = source
            self.output_var.set(output)

        self.run_btn.config(state=tk.DISABLED)
        t = threading.Thread(target=self._run_worker,
                             args=(Path(source), Path(output)),
                             daemon=True)
        t.start()

    def _run_worker(self, source: Path, output: Path):
        try:
            logger.info(f"Scanning: {source}")
            files = find_markdown_files(source)
            if not files:
                logger.warning(f"No *_Synopsis.md files found in {source}")
                messagebox.showwarning(
                    "No files",
                    f"No *_Synopsis.md files found in:\n{source}")
                return

            logger.info(f"Found {len(files)} file(s)")
            for f in files:
                logger.info(f"  - {f.name}")

            ts = datetime.now().strftime("%Y-%m-%d_%H%M%S")
            out_path = output / f"Synopses_{ts}.docx"
            build_document(files, out_path)

            logger.info("Done.")
            if messagebox.askyesno(
                    "Complete",
                    f"Generated:\n{out_path}\n\nOpen it now?"):
                os.startfile(str(out_path))
        except Exception as e:
            logger.exception(f"Failed: {e}")
            messagebox.showerror("Error", f"Generation failed:\n{e}")
        finally:
            self.run_btn.config(state=tk.NORMAL)


# =============================================================================
# MAIN
# =============================================================================

def main():
    root = tk.Tk()
    SynopsisToDocxGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()